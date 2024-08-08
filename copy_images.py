import os
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

total_count = 22

def get_image_paths(folder_path):
    # image 폴더 안의 모든 이미지 파일 경로 가져오기
    image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif')  # 필요한 이미지 확장자들
    image_paths = []
    
    # 폴더 내 파일명 리스트 가져오기 및 정렬
    filenames = sorted(os.listdir(folder_path))
    
    for filename in filenames:
        if filename.lower().endswith(image_extensions):
            image_paths.append(os.path.join(folder_path, filename))
    
    return image_paths

def add_images_inline(doc, front_image_paths, back_image_paths):
    # Word 문서 생성
    doc = Document()

    # 새로운 단락 추가
    paragraph = doc.add_paragraph()

    # 0 1 2 3 4
    for count in range(1, total_count+1):
        # 앞면
        index = count - 1
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run()
        # 이미지를 510mm x 810mm 크기로 설정
        run.add_picture(front_image_paths[index], width=Mm(50), height=Mm(81))
        # 이미지 간격을 위해 공백 추가
        run = paragraph.add_run(" ")
        
        if not count % 10:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for count in range(10):
                run = paragraph.add_run()
                # 이미지를 510mm x 810mm 크기로 설정
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run.add_picture(back_image_paths[0], width=Mm(50), height=Mm(81))
                # 이미지 간격을 위해 공백 추가
                run = paragraph.add_run(" ")

    # Word 문서 저장
    doc.save('images_inline.docx')

# image 폴더 경로
front_image_folder_path = 'front'
back_image_folder_path = 'back'

# 이미지 파일 경로 리스트 생성
front_image_paths = get_image_paths(front_image_folder_path)
back_image_paths = get_image_paths(back_image_folder_path)

# 이미지 일렬로 추가하여 Word 파일로 저장
add_images_inline(Document(), front_image_paths, back_image_paths)
